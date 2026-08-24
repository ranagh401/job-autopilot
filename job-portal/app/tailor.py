"""Tailors the base resume to a specific job and renders a .docx."""
from __future__ import annotations

import re
from pathlib import Path

from . import ats, resume_formats
from .db import DATA_DIR, ResumeVersion
from .llm import chat
from .profile import profile_summary

SYSTEM = (
    "You are an expert technical resume writer. Rewrite the candidate's base "
    "resume so it is tailored to the given job posting: reorder and rephrase "
    "to emphasise the most relevant skills/projects, mirror important "
    "keywords from the job description, and follow the target market's "
    "resume conventions exactly as given.\n"
    "STRICT RULE: never invent employers, job titles, dates, degrees, tools "
    "or metrics that are not in the base resume. You may only reword, "
    "reorder, merge or drop existing content.\n"
    "Respond with JSON exactly in this shape:\n"
    '{"summary": "...", '
    '"skills": [{"category": "...", "items": ["..."]}], '
    '"experience": [{"role": "...", "company": "...", "dates": "...", '
    '"bullets": ["..."]}], '
    '"projects": [{"name": "...", "tech": "...", "bullets": ["..."]}], '
    '"education": [{"degree": "...", "school": "...", "dates": "..."}], '
    '"certifications": ["..."], '
    '"languages": ["e.g. English - C1 (professional)"]}'
)


def extract_text(path: str) -> str:
    p = Path(path)
    if p.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(x for x in parts if x.strip())
    if p.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "")
                         for page in PdfReader(str(p)).pages)
    return p.read_text(encoding="utf-8", errors="ignore")


def _slug(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", text or "x").strip("_")[:40] or "x"


def render_docx(data: dict, profile: dict, out_path: Path,
                fmt: dict | None = None) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(profile.get("name") or "Resume", 0)
    contact_bits = [profile.get("phone"), profile.get("email"),
                    profile.get("current_location")]
    contact_bits += [v for v in (profile.get("links") or {}).values() if v]
    p = doc.add_paragraph(" | ".join(str(b) for b in contact_bits if b))
    p.runs[0].font.size = Pt(9)

    if data.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(data["summary"])
    if data.get("skills"):
        doc.add_heading("Skills", level=2)
        for s in data["skills"]:
            para = doc.add_paragraph()
            para.add_run(f"{s.get('category', 'Skills')}: ").bold = True
            para.add_run(", ".join(s.get("items") or []))
    if data.get("experience"):
        doc.add_heading("Experience", level=2)
        for e in data["experience"]:
            para = doc.add_paragraph()
            para.add_run(f"{e.get('role', '')} — {e.get('company', '')}").bold = True
            if e.get("dates"):
                para.add_run(f"  ({e['dates']})")
            for b in e.get("bullets") or []:
                doc.add_paragraph(b, style="List Bullet")
    if data.get("projects"):
        doc.add_heading("Projects", level=2)
        for pr in data["projects"]:
            para = doc.add_paragraph()
            para.add_run(pr.get("name", "")).bold = True
            if pr.get("tech"):
                para.add_run(f"  [{pr['tech']}]")
            for b in pr.get("bullets") or []:
                doc.add_paragraph(b, style="List Bullet")
    if data.get("education"):
        doc.add_heading("Education", level=2)
        for ed in data["education"]:
            doc.add_paragraph(
                f"{ed.get('degree', '')} — {ed.get('school', '')} "
                f"{('(' + ed['dates'] + ')') if ed.get('dates') else ''}")
    if data.get("certifications"):
        doc.add_heading("Certifications", level=2)
        for c in data["certifications"]:
            doc.add_paragraph(c, style="List Bullet")
    # Germany and the Netherlands expect a Languages section with levels.
    if data.get("languages"):
        doc.add_heading("Languages", level=2)
        doc.add_paragraph(", ".join(str(l) for l in data["languages"]))

    doc.save(str(out_path))
    return out_path


def _resume_text(data: dict) -> str:
    """Flatten the generated resume JSON so it can be ATS-scored."""
    parts = [data.get("summary", "")]
    for s in data.get("skills") or []:
        parts.append(s.get("category", ""))
        parts += s.get("items") or []
    for e in data.get("experience") or []:
        parts += [e.get("role", ""), e.get("company", "")]
        parts += e.get("bullets") or []
    for p in data.get("projects") or []:
        parts += [p.get("name", ""), p.get("tech", "")]
        parts += p.get("bullets") or []
    for ed in data.get("education") or []:
        parts += [ed.get("degree", ""), ed.get("school", "")]
    parts += data.get("certifications") or []
    # Section headings the renderer emits, so completeness scores honestly
    parts += ["Summary", "Skills", "Experience", "Education"]
    return "\n".join(str(p) for p in parts if p)


def tailor_resume(session, job, profile) -> dict:
    base = profile.get("base_resume")
    if not base:
        raise RuntimeError("No base resume uploaded yet - "
                           "upload it from the dashboard first.")
    base_text = extract_text(base)
    desc = (job.description or "")[:6000]
    reqs = ats.extract_requirements(desc)
    want = ", ".join((reqs["required_skills"] + reqs["keywords"])[:20])

    # A resume for Germany or Australia is not the same document as one
    # for India - follow the destination market's conventions.
    fmt = resume_formats.for_country(job.country, job.is_india)
    abroad = not (job.is_india if job.is_india is not None else True)
    needs_visa = abroad and bool(profile.get("needs_visa_sponsorship"))
    user = (
        "CANDIDATE PROFILE:\n" + profile_summary(profile) + "\n\n"
        "BASE RESUME:\n" + base_text[:9000] + "\n\n"
        "JOB POSTING:\n"
        f"Title: {job.title}\nCompany: {job.company}\n"
        f"Location: {job.location} ({job.country or 'unknown country'})\n"
        f"Description:\n{desc}\n\n"
        "RESUME FORMAT RULES FOR THIS MARKET:\n"
        + resume_formats.prompt_block(fmt, needs_visa) + "\n\n"
        f"ATS KEYWORDS this posting screens for: {want}\n"
        "Where the base resume genuinely demonstrates one of these, use the "
        "posting's exact wording for it. Never claim one it does not support."
    )
    data = chat(SYSTEM, user, json_mode=True, kind="tailor-resume")

    # One improvement pass if important keywords the candidate plausibly
    # has are still absent (Resume-Matcher's keyword-injection idea).
    first = ats.score_resume(_resume_text(data), desc, reqs)
    if first["missing"] and first["overall"] < 85:
        retry = chat(SYSTEM, user + (
            "\n\nA first draft missed these keywords: "
            + ", ".join(first["missing"])
            + ". Rewrite it working in ONLY those the base resume genuinely "
              "supports (as evidence-backed phrasing, never a bare keyword "
              "list). Ignore any the candidate cannot honestly claim."),
            json_mode=True)
        if ats.score_resume(_resume_text(retry), desc, reqs)["overall"] > \
                first["overall"]:
            data = retry

    suffix = _slug(fmt["label"]) if abroad else "IN"
    out_path = (DATA_DIR / "resumes" /
                f"job{job.id}_{_slug(job.company)}_{suffix}.docx")
    render_docx(data, profile, out_path, fmt)
    rv = ResumeVersion(job_id=job.id, path=str(out_path))
    session.add(rv)
    session.commit()
    result = ats.score_and_store(session, rv, _resume_text(data), desc, reqs)
    if job.status in ("found", "scored", "shortlisted"):
        job.status = "tailored"
    session.commit()
    return {"path": str(out_path), "ats": result, "format": fmt["label"],
            "doc_name": fmt["doc_name"]}
